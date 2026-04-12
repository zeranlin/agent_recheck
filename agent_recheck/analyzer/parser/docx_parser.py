"""DOCX 文档解析器"""

import re
from pathlib import Path
from typing import Optional, List, Any
from dataclasses import dataclass, field
from datetime import datetime


@dataclass
class DocumentMetadata:
    """文档元数据"""
    file_path: str = ""
    file_name: str = ""
    file_size: int = 0
    file_type: str = "docx"
    paragraph_count: int = 0
    table_count: int = 0


@dataclass
class DocumentSection:
    """文档章节"""
    title: str = ""
    level: int = 1
    start_line: int = 0
    end_line: int = 0
    content: str = ""


@dataclass
class TableInfo:
    """表格信息"""
    index: int = 0
    title: Optional[str] = None
    rows: int = 0
    cols: int = 0
    start_line: int = 0
    end_line: int = 0
    is_nested: bool = False


@dataclass
class MarkedContent:
    """标记内容"""
    type: str = ""  # 实质性/重要参数
    content: str = ""
    line: int = 0
    context: str = ""


@dataclass
class Document:
    """文档对象"""
    metadata: Optional[DocumentMetadata] = None
    full_text: str = ""
    sections: List[DocumentSection] = field(default_factory=list)
    tables: List[TableInfo] = field(default_factory=list)
    marked_contents: List[MarkedContent] = field(default_factory=list)
    paragraphs: List[str] = field(default_factory=list)
    parsed_at: datetime = field(default_factory=datetime.now)


class DocxParser:
    """DOCX 文档解析器"""

    def __init__(self):
        pass

    def parse(self, file_path: Path) -> Document:
        """解析 DOCX 文件"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument(file_path)

        metadata = self._extract_metadata(file_path, doc)
        full_text = self.extract_text(file_path)
        sections = self._extract_sections(doc)
        tables = self._extract_tables(doc)
        paragraphs = self._extract_paragraphs(doc)
        marked_contents = self._extract_marked_contents(doc)

        return Document(
            metadata=metadata,
            full_text=full_text,
            sections=sections,
            tables=tables,
            marked_contents=marked_contents,
            paragraphs=paragraphs,
        )

    def extract_text(self, file_path: Path) -> str:
        """提取纯文本"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        return "\n".join(paragraphs)

    def extract_tables(self, file_path: Path) -> List[dict]:
        """提取表格"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument(file_path)
        tables = []

        for i, table in enumerate(doc.tables):
            table_data = self._parse_table(table)
            table_data["index"] = i
            tables.append(table_data)

        return tables

    def _extract_metadata(self, file_path: Path, doc) -> DocumentMetadata:
        """提取元数据"""
        return DocumentMetadata(
            file_path=str(file_path),
            file_name=file_path.name,
            file_size=file_path.stat().st_size if file_path.exists() else 0,
            file_type="docx",
            paragraph_count=len([p for p in doc.paragraphs if p.text.strip()]),
            table_count=len(doc.tables),
        )

    def _extract_sections(self, doc) -> List[DocumentSection]:
        """提取章节结构"""
        sections = []
        current_line = 0

        chapter_pattern = re.compile(r"^第[一二三四五六七八九十百\d]+[章节条]")
        section_pattern = re.compile(r"^[一二三四五六七八九十]+、")
        article_pattern = re.compile(r"^\d+[.、]")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                current_line += 1
                continue

            level = 0
            if chapter_pattern.match(text):
                level = 1
            elif section_pattern.match(text):
                level = 2
            elif article_pattern.match(text):
                level = 3

            if level > 0:
                sections.append(
                    DocumentSection(
                        title=text[:80],
                        level=level,
                        start_line=current_line,
                        end_line=current_line,
                        content=text,
                    )
                )

            current_line += 1

        return sections

    def _extract_tables(self, doc) -> List[TableInfo]:
        """提取表格信息"""
        tables = []

        for i, table in enumerate(doc.tables):
            table_info = self._parse_table_info(table, i)
            if table_info:
                tables.append(table_info)

        return tables

    def _extract_paragraphs(self, doc) -> List[str]:
        """提取段落"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return paragraphs

    def _parse_table_info(self, table, index: int) -> Optional[TableInfo]:
        """解析表格信息"""
        rows = len(table.rows)
        cols = len(table.columns) if rows > 0 else 0
        is_nested = rows < 2 or cols < 2

        title = None
        if rows > 0:
            first_cell = table.cell(0, 0).text.strip()
            if first_cell:
                title = first_cell[:50]

        return TableInfo(
            index=index,
            title=title,
            rows=rows,
            cols=cols,
            start_line=0,
            end_line=0,
            is_nested=is_nested,
        )

    def _extract_marked_contents(self, doc) -> List[MarkedContent]:
        """提取标记内容"""
        marked = []
        line = 0

        for para in doc.paragraphs:
            text = para.text
            line += 1

            if "★" in text:
                marked.append(
                    MarkedContent(
                        type="实质性",
                        content=text.strip(),
                        line=line,
                        context=text.strip()[:200],
                    )
                )

            if "▲" in text:
                marked.append(
                    MarkedContent(
                        type="重要参数",
                        content=text.strip(),
                        line=line,
                        context=text.strip()[:200],
                    )
                )

        return marked

    def _parse_table(self, table) -> dict:
        """解析表格内容"""
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        return {"data": data}
