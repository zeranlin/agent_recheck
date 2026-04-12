# -*- coding: utf-8 -*-
"""
解析器测试

测试文档解析功能
"""

import pytest
import tempfile
import os
from pathlib import Path

# 导入被测试模块
from analyzer.parser.docx_parser import DocxParser
from analyzer.parser.base import ParsedDocument


class TestDocxParser:
    """DOCX 解析器测试"""
    
    @pytest.fixture
    def sample_docx_path(self):
        """创建示例 DOCX 文件"""
        # 在临时目录创建测试文件
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            
            # 添加标题
            doc.add_heading('政府采购招标文件', 0)
            doc.add_paragraph('这是一份测试文件')
            
            # 添加实质性条款标记
            doc.add_paragraph('★ 本条款为实质性条款')
            
            # 添加表格
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # 填写表格内容
            cells = table.rows[0].cells
            cells[0].text = '评分因素'
            cells[1].text = '分值'
            cells[2].text = '评分标准'
            
            doc.save(f.name)
            
            yield f.name
            
            # 清理
            os.unlink(f.name)
    
    def test_parser_initialization(self):
        """测试解析器初始化"""
        parser = DocxParser()
        assert parser is not None
    
    def test_parse_document(self, sample_docx_path):
        """测试文档解析"""
        parser = DocxParser()
        doc = parser.parse(sample_docx_path)
        
        assert doc is not None
        assert isinstance(doc, ParsedDocument)
        assert doc.file_name is not None
    
    def test_extract_text(self, sample_docx_path):
        """测试文本提取"""
        parser = DocxParser()
        doc = parser.parse(sample_docx_path)
        
        # 检查是否包含预期内容
        full_text = doc.full_text
        assert '政府采购招标文件' in full_text
        assert '测试文件' in full_text
    
    def test_extract_tables(self, sample_docx_path):
        """测试表格提取"""
        parser = DocxParser()
        doc = parser.parse(sample_docx_path)
        
        # 检查表格数量
        assert len(doc.tables) > 0
    
    def test_marked_content(self, sample_docx_path):
        """测试标记内容提取"""
        parser = DocxParser()
        doc = parser.parse(sample_docx_path)
        
        # 检查是否提取了标记内容
        assert '★' in doc.full_text


class TestTableParser:
    """表格解析器测试"""
    
    def test_table_type_detection(self):
        """测试表格类型检测"""
        from analyzer.parser.enhanced_table_parser import EnhancedTableParser, TableType
        
        parser = EnhancedTableParser()
        
        # 创建测试表格数据
        mock_rows = []
        
        # 测试评分标准表识别
        mock_text = "评分因素 价格分 技术分"
        detected = parser._identify_table_type(mock_rows)
        assert detected == TableType.UNKNOWN  # 空数据
    
    def test_cell_classification(self):
        """测试单元格分类"""
        from analyzer.parser.enhanced_table_parser import EnhancedTableParser, CellType
        
        parser = EnhancedTableParser()
        
        # 测试不同类型的单元格
        assert parser._classify_cell("★ 实质性条款", False) == CellType.HIGHLIGHTED_STAR
        assert parser._classify_cell("▲ 重要参数", False) == CellType.HIGHLIGHTED_TRIANGLE
        assert parser._classify_cell("普通内容", False) == CellType.NORMAL
        assert parser._classify_cell("表头", True) == CellType.HEADER


class TestScoringParser:
    """评分解析器测试"""
    
    def test_scoring_standard_parsing(self):
        """测试评分标准解析"""
        from analyzer.engine.scoring_parser import ScoringParser, EvaluationMethod
        
        parser = ScoringParser()
        
        # 测试评审方法检测
        text_lowest = "本项目采用最低价法评审"
        method = parser._detect_evaluation_method(text_lowest)
        assert method == EvaluationMethod.LOWEST_PRICE
        
        text_comprehensive = "本项目采用综合评分法"
        method = parser._detect_evaluation_method(text_comprehensive)
        assert method == EvaluationMethod.COMPREHENSIVE
    
    def test_price_weight_extraction(self):
        """测试价格分权重提取"""
        from analyzer.engine.scoring_parser import ScoringParser
        
        parser = ScoringParser()
        
        # 测试各种格式的价格权重提取
        test_cases = [
            ("价格分权重: 30%", 30.0),
            ("价格分: 50%", 50.0),
            ("价格 60%", 60.0),
        ]
        
        for text, expected in test_cases:
            weight = parser._extract_price_weight(text)
            assert weight == expected, f"Failed for {text}"


class TestShenzhenAdapter:
    """深圳格式适配器测试"""
    
    def test_adapter_initialization(self):
        """测试适配器初始化"""
        from analyzer.parser.shenzhen_adapter import ShenzhenAdapter
        
        adapter = ShenzhenAdapter()
        assert adapter is not None
    
    def test_region_detection(self):
        """测试地区检测"""
        from analyzer.parser.shenzhen_adapter import ShenzhenAdapter
        
        adapter = ShenzhenAdapter()
        
        # 测试深圳地区检测
        text = "深圳市政府采购招标文件 SZCG20240001"
        region = adapter.detect_region(text)
        assert region == "深圳"
        
        # 测试广东地区检测
        text = "广东省政府采购 GD20240001"
        region = adapter.detect_region(text)
        assert region == "广东"


# 运行测试
if __name__ == "__main__":
    pytest.main([__file__, "-v"])
